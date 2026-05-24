"""
학생부 종합 면접 어시스트 v3.0
- 학년/영역 연계형 질문과 단일 활동 심화형(꼬리질문) 균형 생성 로직 반영
- 세트당 다수의 꼬리/연계질문(2~4개) 강제 할당 로직 추가 (심층면접 강화)
- 화면 표시 / DOCX 다운로드 지원 
- 하단 고정 푸터(만든 이) 유지
"""

import streamlit as st
import json
import io
from datetime import datetime
from pathlib import Path

from google import genai
from google.genai import types

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pypdf

# ═══════════════════════════════════════════════════════════
# 페이지 설정
# ═══════════════════════════════════════════════════════════
st.set_page_config(
    page_title="학생부 면접 질문 생성기 시스템 v3.0",
    page_icon="🎤",
    layout="wide",
)

# ═══════════════════════════════════════════════════════════
# 하단 푸터 CSS (중앙 정렬 및 스타일링)
# ═══════════════════════════════════════════════════════════
footer_css = """
<style>
.footer {
    position: fixed;
    left: 0;
    bottom: 0;
    width: 100%;
    background-color: white;
    color: #555555;
    text-align: center;
    padding: 15px 0;
    font-size: 14px;
    border-top: 1px solid #e5e7eb;
    z-index: 1000;
    line-height: 1.6;
}
/* 푸터에 가리지 않도록 메인 컨테이너 하단 여백 추가 */
.main .block-container {
    padding-bottom: 100px; 
}
</style>
<div class="footer">
    <b>🏫 학생부 면접 질문 생성기 시스템 v3.0</b><br>
    만든 이: 신선여자고등학교 김명남<br>
    🗓️ 2026.05
</div>
"""
st.markdown(footer_css, unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════
# 시스템 프롬프트 (다중 서브질문 강제)
# ═══════════════════════════════════════════════════════════
SYSTEM_PROMPT = """
당신은 대학 입학사정관 및 면접 출제 전문가입니다. 제공된 학생부 내용을 분석하여, '연계형' 질문과 '심화형(꼬리질문)' 질문을 균형 있게 섞어서 심층 면접 질문지를 구성하세요.

[질문 유형 정의]
1. 연계형 (Linked): 특정 학년/과목의 관심사가 다른 학년/과목으로 확장된 지점을 찾아 연결하여 묻는 다각도 질문.
2. 심화형 (Deep-dive): 하나의 주요 활동에 대해 동기(WHY), 과정(HOW), 결과 및 배움(LEARN)을 깊이 있게 연속으로 파고드는 압박형 꼬리질문.

[출제 절대 원칙 - 반드시 지킬 것]
1. 사용자가 요청한 '메인 질문 세트 수'에 맞게 두 유형(연계형, 심화형)을 자연스럽게 혼합하세요.
2. **핵심 원칙**: 하나의 메인 질문 아래에는 사용자가 요청한 '세트당 연계/꼬리질문 수'만큼 반드시 서브 질문(sub_list)을 2개 이상 꽉 채워서 생성해야 합니다. (예: 1-1, 1-2, 1-3 연속 출제)
3. 질문 끝에는 반드시 기재된 출처(예: 1학년 통합사회, 2학년 진로활동)를 괄호로 명시하세요.
4. 각 서브 질문에 대해 학생부에 기록된 사실을 바탕으로 '모범 답변(또는 활동 내용 요약)'을 제시하세요.

[필수 JSON 출력 형식]
반드시 아래의 JSON 구조로만 출력하세요.
{
  "questions": [
    {
      "no": 1,
      "question_type": "연계형",
      "main_question": "유전공학에 대한 관심으로 유전자 조작 유기체 문제를 조사하고 발표함, 어떤 내용인가요?",
      "main_context": "(1학년 통합사회)",
      "sub_list": [
        {
          "sub_no": "1-1",
          "question": "첫 번째 연계/꼬리질문 내용",
          "context": "(2학년 생명과학 I)",
          "expected_answer": "기대 답변 1"
        },
        {
          "sub_no": "1-2",
          "question": "두 번째 연계/꼬리질문 내용",
          "context": "(2학년 생명과학 I)",
          "expected_answer": "기대 답변 2"
        },
        {
          "sub_no": "1-3",
          "question": "세 번째 연계/꼬리질문 내용 (사용자가 3개를 요청했을 경우)",
          "context": "(3학년 진로활동)",
          "expected_answer": "기대 답변 3"
        }
      ]
    }
  ]
}
"""

# ═══════════════════════════════════════════════════════════
# 유틸리티 함수
# ═══════════════════════════════════════════════════════════
def extract_text_from_pdf(file) -> str:
    """PDF 텍스트 추출"""
    try:
        reader = pypdf.PdfReader(file)
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as e:
        st.error(f"PDF 읽기 실패: {e}")
        return ""

def call_gemini(api_key: str, model: str, user_input: str) -> dict:
    """Gemini API 호출 및 JSON 파싱"""
    client = genai.Client(api_key=api_key)
    response = client.models.generate_content(
        model=model,
        contents=[{"role": "user", "parts": [{"text": user_input}]}],
        config=types.GenerateContentConfig(
            system_instruction=SYSTEM_PROMPT,
            temperature=0.4,
            response_mime_type="application/json",
        ),
    )
    raw = response.text.strip()
    raw = raw.replace("```json", "").replace("```", "").strip()
    return json.loads(raw)

def build_docx(sheet: dict, meta: dict) -> bytes:
    """질문지를 DOCX 파일로 변환"""
    doc = Document()
    
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(11)

    title = doc.add_heading("모의 면접 질문지", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run(f"희망 전공: {meta['major']} | 생성일: {meta['date']} | 메인 세트 {len(sheet['questions'])}개").italic = True
    doc.add_paragraph("─" * 50)

    for q in sheet["questions"]:
        # 메인 질문
        mp = doc.add_paragraph()
        run = mp.add_run(f"{q['no']}. {q.get('main_question', '')} {q.get('main_context', '')}")
        run.font.size = Pt(12)

        # 서브 질문 (유형에 따라 타이틀 변경)
        if q.get("sub_list"):
            q_type = q.get("question_type", "심화형")
            label = "[연계질문]" if q_type == "연계형" else "[꼬리질문]"
            doc.add_paragraph(label).bold = True
            
            for sub in q["sub_list"]:
                # 질문
                lp = doc.add_paragraph()
                lp.add_run(f"{sub.get('sub_no', '')}. {sub.get('question', '')} {sub.get('context', '')}")
                
                # 기대 답변
                ap = doc.add_paragraph()
                ap.add_run(f": {sub.get('expected_answer', '')}")
                
        doc.add_paragraph() # 문항 간 간격

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# ═══════════════════════════════════════════════════════════
# 사이드바: 설정
# ═══════════════════════════════════════════════════════════
with st.sidebar:
    st.markdown("### 🔑 기본 설정")
    api_key = st.text_input("Google AI API 키", type="password")
    
    st.divider()
    st.markdown("### ⚙️ 모델 설정")
    model_name = st.selectbox(
        "Gemini 모델",
        ["gemini-2.5-flash", "gemini-2.5-pro"],
        index=0,
    )

    st.divider()
    st.markdown("### 📊 질문지 옵션")
    n_main = st.slider("메인 질문 세트 수", 3, 12, 6)
    n_tail = st.slider("세트당 연계/꼬리질문 수", 2, 4, 2, help="하나의 메인 질문에 달리는 추가 심화 질문의 개수입니다.")

# ═══════════════════════════════════════════════════════════
# 메인 레이아웃
# ═══════════════════════════════════════════════════════════
st.title("🎤 학생부 면접 질문 생성기 시스템 v3.0")
st.caption("학생부의 맥락을 파악하여 '융합형 연계질문'과 '단일활동 심화 꼬리질문'을 심층적으로 생성합니다.")

col_left, col_right = st.columns([1.2, 1])

with col_left:
    st.subheader("1. 학생부 데이터 입력")
    uploaded = st.file_uploader("📁 학생부 PDF 업로드 (선택)", type=["pdf"])
    
    extracted = ""
    if uploaded is not None:
        with st.spinner("PDF에서 텍스트 추출 중..."):
            extracted = extract_text_from_pdf(uploaded)
        if extracted.strip():
            st.success(f"PDF 추출 완료!")

    record_text = st.text_area(
        "📝 학생부 내용 (직접 붙여넣기 권장)",
        value=extracted, 
        height=400,
        placeholder="[1학년 통합사회] ...\n[2학년 동아리활동] ..."
    )

with col_right:
    st.subheader("2. 분석 옵션")
    major = st.text_input("🎓 희망 전공 / 학과", placeholder="예: 생명공학과")
    
    st.markdown("---")
    st.markdown("**🏫 목표 대학 면접 스타일 (선택)**")
    univ_uploaded = st.file_uploader("대학별 면접 가이드북 / 대학별 면접 기출 문제 업로드", type=["pdf"])
    
    st.divider()
    run = st.button("🚀 심층 면접 질문지 생성", type="primary", use_container_width=True)

# ═══════════════════════════════════════════════════════════
# 실행부
# ═══════════════════════════════════════════════════════════
if run:
    if not api_key:
        st.error("Google AI API 키를 입력해주세요.")
        st.stop()
    if not record_text.strip():
        st.error("학생부 내용을 입력해주세요.")
        st.stop()
    if not major.strip():
        st.error("희망 전공을 입력해주세요.")
        st.stop()

    # 프롬프트 세팅
    user_input = f"희망 전공: {major}\n생성할 전체 메인 질문 세트 수: {n_main}\n세트당 반드시 생성할 연계/꼬리질문 수: {n_tail}개\n\n"
    
    univ_guide_text = ""
    if univ_uploaded is not None:
        univ_guide_text = extract_text_from_pdf(univ_uploaded)
        if univ_guide_text.strip():
            user_input += f"[대학 가이드북 참고]\n{univ_guide_text}\n\n"

    user_input += f"[학생부 원문]\n{record_text}"

    # AI 생성
    with st.spinner(f"학생부 맥락을 분석하여 세트당 {n_tail}개의 심층 질문을 추출 중입니다..."):
        try:
            sheet = call_gemini(api_key, model_name, user_input)
        except Exception as e:
            st.error(f"생성 실패: {e}")
            st.stop()

    questions = sheet.get("questions", [])
    if not questions:
        st.error("질문이 생성되지 않았습니다.")
        st.stop()

    st.success(f"✅ 총 {len(questions)}개의 면접 질문 세트를 생성했습니다.")

    # ────────────────────────────────────────────────
    # 화면 출력
    # ────────────────────────────────────────────────
    st.markdown("---")
    st.markdown("## 📋 생성된 심층 면접 질문지")
    
    for q in questions:
        with st.container(border=True):
            # 메인 질문
            st.markdown(f"#### {q.get('no', '')}. {q.get('main_question', '')} **{q.get('main_context', '')}**")
            
            # 서브 질문 (유형에 따라 타이틀 분기)
            if q.get("sub_list"):
                q_type = q.get("question_type", "심화형")
                label = "[연계질문]" if q_type == "연계형" else "[꼬리질문]"
                
                st.markdown(f"**{label}**")
                for sub in q["sub_list"]:
                    st.markdown(f"{sub.get('sub_no', '')}. {sub.get('question', '')} **{sub.get('context', '')}**")
                    st.caption(f": {sub.get('expected_answer', '')}")

    # ────────────────────────────────────────────────
    # 다운로드 버튼
    # ────────────────────────────────────────────────
    st.markdown("---")
    st.markdown("### 📥 다운로드")

    meta = {
        "major": major,
        "date": datetime.now().strftime("%Y-%m-%d"),
    }
    base_name = f"면접질문지_{major}_{datetime.now().strftime('%H%M')}"

    docx_bytes = build_docx(sheet, meta)
    st.download_button(
        "📝 Word(docx) 다운로드", 
        data=docx_bytes, 
        file_name=f"{base_name}.docx", 
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
        use_container_width=True
    )
